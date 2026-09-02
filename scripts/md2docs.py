#!/usr/bin/env python3
"""
Convertit un document Markdown en Word (.docx) et en PDF mis en page avec le logo.

Usage :
    python3 scripts/md2docs.py docs/CDC-VOOMNET.md
    python3 scripts/md2docs.py docs/CDC-VOOMNET.md --titre "Cahier des charges" --sous-titre "VOOMNET TECHNOLOGY"
Produit : <fichier>.docx et <fichier>.pdf à côté du fichier source.
"""
import os
import re
import sys
import argparse

VIOLET = "#500070"
BLEU_NUIT = "#000060"
GRIS = "#475569"

# ----------------------------------------------------------------- lecture ---
def nettoyer(texte, pour_pdf=False):
    """Remplace les pictogrammes par du texte et supprime les caractères non latin-1."""
    if not texte:
        return ""
    substitutions = {
        "✅": "[OK]", "⛔": "[KO]", "❌": "[X]", "✔": "[OK]", "⚠️": "[!]", "⚠": "[!]",
        "🔴": "(urgent)", "🟠": "(alerte)", "🔵": "(info)",
        "🔔": "(cloche)", "🔊": "(son)", "🔇": "(muet)", "🔕": "(muet)",
        "📈": "(graphique)", "📊": "(excel)", "📄": "(document)", "🖨️": "(imprimer)",
        "📝": "(demande)", "🏢": "(fournisseurs)", "📦": "(commande)", "🚚": "(reception)",
        "👥": "(utilisateurs)", "⚙️": "(parametres)", "🗑️": "(supprimer)", "🏅": "(top)",
        "💰": "(prix)", "📋": "(liste)", "💳": "(paiement)", "🛡️": "(garantie)",
        "⏳": "(attente)", "🍩": "(repartition)", "🕒": "(recent)", "🛒": "(achats)",
        "→": "->", "←": "<-", "⇄": "<->", "▲": "^", "▼": "v",
        "│": "|", "├": "+", "└": "+", "─": "-", "✔️": "[OK]", "’": "'",
    }
    for cle, valeur in substitutions.items():
        texte = texte.replace(cle, valeur)
    if pour_pdf:
        # on ne conserve que les caractères imprimables par les polices standard
        resultat = []
        for c in texte:
            try:
                c.encode("cp1252")
                resultat.append(c)
            except UnicodeEncodeError:
                resultat.append("")
        texte = "".join(resultat)
    return re.sub(r" {2,}", " ", texte).strip()


def entetes_lignes(lignes):
    """Découpe les lignes d'un tableau Markdown."""
    lignes = [l for l in lignes if not re.match(r"^\|?[\s:|-]+\|?$", l)]
    return [[c.strip() for c in l.strip().strip("|").split("|")] for l in lignes]


def lire_markdown(chemin):
    with open(chemin, encoding="utf-8") as f:
        lignes = f.read().split("\n")
    blocs, i = [], 0
    while i < len(lignes):
        l = lignes[i]
        if l.startswith("```"):
            i += 1
            tampon = []
            while i < len(lignes) and not lignes[i].startswith("```"):
                tampon.append(lignes[i])
                i += 1
            i += 1
            blocs.append(("code", "\n".join(tampon)))
            continue
        if re.match(r"^\s*\|", l):
            tampon = []
            while i < len(lignes) and re.match(r"^\s*\|", lignes[i]):
                tampon.append(lignes[i])
                i += 1
            blocs.append(("table", entetes_lignes(tampon)))
            continue
        if l.startswith("### "):
            blocs.append(("h3", l[4:].strip()))
            i += 1
            continue
        if l.startswith("## "):
            blocs.append(("h2", l[3:].strip()))
            i += 1
            continue
        if l.startswith("# "):
            blocs.append(("h1", l[2:].strip()))
            i += 1
            continue
        if re.match(r"^\s*[-*] ", l):
            tampon = []
            while i < len(lignes) and re.match(r"^\s*[-*] ", lignes[i]):
                tampon.append(re.sub(r"^\s*[-*] ", "", lignes[i]).strip())
                i += 1
            blocs.append(("ul", tampon))
            continue
        if re.match(r"^\s*\d+\. ", l):
            tampon = []
            while i < len(lignes) and re.match(r"^\s*\d+\. ", lignes[i]):
                tampon.append(re.sub(r"^\s*\d+\. ", "", lignes[i]).strip())
                i += 1
            blocs.append(("ol", tampon))
            continue
        if l.startswith("> "):
            blocs.append(("quote", l[2:].strip()))
            i += 1
            continue
        if re.match(r"^\s*---+\s*$", l):
            blocs.append(("hr", ""))
            i += 1
            continue
        if not l.strip():
            i += 1
            continue
        tampon = []
        while (i < len(lignes) and lignes[i].strip()
               and not re.match(r"^(#{1,3} |\s*\| |\s*[-*] |\s*\d+\. |> |```)", lignes[i])):
            tampon.append(lignes[i].strip())
            i += 1
        if tampon:
            blocs.append(("p", " ".join(tampon)))
    return blocs


def morceaux_gras(texte):
    """Découpe un texte en (morceau, est_gras)."""
    return [(m, bool(m.startswith("**") and m.endswith("**")))
            for m in re.split(r"(\*\*[^*]+\*\*)", texte) if m]


# ------------------------------------------------------------------- DOCX ---
def generer_docx(blocs, sortie, titre, sous_titre):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for section in doc.sections:
        section.left_margin = section.right_margin = Cm(2)
        section.top_margin = section.bottom_margin = Cm(2)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    # page de garde
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(titre)
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor(0x50, 0x00, 0x70)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(sous_titre)
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x00, 0x00, 0x60)
    doc.add_paragraph()
    logo = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "public", "voomnet-logo.png")
    if os.path.exists(logo):
        doc.add_picture(logo, width=Cm(9))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    for type_bloc, contenu in blocs:
        if type_bloc == "h1":
            h = doc.add_heading(nettoyer(contenu), 0)
        elif type_bloc == "h2":
            doc.add_heading(nettoyer(contenu), 1)
        elif type_bloc == "h3":
            doc.add_heading(nettoyer(contenu), 2)
        elif type_bloc == "p":
            para = doc.add_paragraph()
            for morceau, gras in morceaux_gras(nettoyer(contenu)):
                run = para.add_run(morceau.strip("*"))
                run.bold = gras
        elif type_bloc == "ul":
            for item in contenu:
                para = doc.add_paragraph(style="List Bullet")
                for morceau, gras in morceaux_gras(nettoyer(item)):
                    run = para.add_run(morceau.strip("*"))
                    run.bold = gras
        elif type_bloc == "ol":
            for item in contenu:
                para = doc.add_paragraph(style="List Number")
                for morceau, gras in morceaux_gras(nettoyer(item)):
                    run = para.add_run(morceau.strip("*"))
                    run.bold = gras
        elif type_bloc == "quote":
            para = doc.add_paragraph()
            run = para.add_run(nettoyer(contenu))
            run.italic = True
            run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
        elif type_bloc == "code":
            para = doc.add_paragraph()
            run = para.add_run(contenu)
            run.font.name = "Consolas"
            run.font.size = Pt(8.5)
        elif type_bloc == "table" and contenu:
            nb_col = max(len(ligne) for ligne in contenu)
            tableau = doc.add_table(rows=len(contenu), cols=nb_col)
            tableau.style = "Table Grid"
            for y, ligne in enumerate(contenu):
                for x in range(nb_col):
                    cellule = tableau.cell(y, x)
                    cellule.text = ""
                    para = cellule.paragraphs[0]
                    texte = nettoyer(ligne[x]) if x < len(ligne) else ""
                    run = para.add_run(texte)
                    if y == 0:
                        run.bold = True
                    run.font.size = Pt(8.5)
            doc.add_paragraph()

    doc.save(sortie)
    return sortie


# -------------------------------------------------------------------- PDF ---
def generer_pdf(blocs, sortie, titre, sous_titre):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import (BaseDocTemplate, PageTemplate, Frame, Paragraph, Spacer,
                                    Table as PdfTable, TableStyle, PageBreak, Image, KeepTogether)

    violet = colors.HexColor(VIOLET)
    bleu = colors.HexColor(BLEU_NUIT)
    gris = colors.HexColor(GRIS)

    styles = {
        "h1": ParagraphStyle("h1", fontName="Helvetica-Bold", fontSize=16, textColor=violet,
                             spaceBefore=14, spaceAfter=8),
        "h2": ParagraphStyle("h2", fontName="Helvetica-Bold", fontSize=12.5, textColor=bleu,
                             spaceBefore=12, spaceAfter=6),
        "h3": ParagraphStyle("h3", fontName="Helvetica-Bold", fontSize=10.5, textColor=violet,
                             spaceBefore=9, spaceAfter=4),
        "p": ParagraphStyle("p", fontName="Helvetica", fontSize=9, leading=12.5,
                            textColor=colors.HexColor("#0f172a"), spaceAfter=5),
        "ul": ParagraphStyle("ul", fontName="Helvetica", fontSize=9, leading=12.5,
                             leftIndent=10, bulletIndent=2, spaceAfter=2),
        "quote": ParagraphStyle("quote", fontName="Helvetica-Oblique", fontSize=8.5, leading=11.5,
                                textColor=gris, leftIndent=10, spaceAfter=5),
        "code": ParagraphStyle("code", fontName="Courier", fontSize=7.6, leading=9.6,
                               backColor=colors.HexColor("#f1f5f9"), leftIndent=6, spaceAfter=6),
        "cell": ParagraphStyle("cell", fontName="Helvetica", fontSize=7.6, leading=9.6),
        "cellh": ParagraphStyle("cellh", fontName="Helvetica-Bold", fontSize=7.6, leading=9.6,
                                textColor=colors.white),
        "titre": ParagraphStyle("titre", fontName="Helvetica-Bold", fontSize=24,
                                textColor=violet, alignment=1, spaceAfter=10),
        "sous": ParagraphStyle("sous", fontName="Helvetica", fontSize=13,
                               textColor=bleu, alignment=1, spaceAfter=4),
    }

    def enrichir(texte):
        t = nettoyer(texte, pour_pdf=True)
        t = t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        t = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", t)
        t = re.sub(r"`([^`]+)`", r'<font face="Courier" size="8">\1</font>', t)
        return t

    recit = []
    logo = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "public", "voomnet-logo.png")
    if os.path.exists(logo):
        recit.append(Spacer(1, 2 * cm))
        recit.append(Image(logo, width=11 * cm, height=11 * cm * 600 / 1920))
        recit.append(Spacer(1, 1 * cm))
    recit.append(Paragraph(enrichir(titre), styles["titre"]))
    recit.append(Paragraph(enrichir(sous_titre), styles["sous"]))
    recit.append(Spacer(1, 1.5 * cm))
    recit.append(Paragraph(
        '<para alignment="center"><font size="9" color="#64748b">Document généré le %s</font></para>'
        % __import__("datetime").date.today().strftime("%d/%m/%Y"), styles["p"]))
    recit.append(PageBreak())

    for type_bloc, contenu in blocs:
        if type_bloc in ("h1", "h2", "h3"):
            recit.append(Paragraph(enrichir(contenu), styles[type_bloc]))
        elif type_bloc == "p":
            recit.append(Paragraph(enrichir(contenu), styles["p"]))
        elif type_bloc == "ul":
            for item in contenu:
                recit.append(Paragraph("&bull; " + enrichir(item), styles["ul"]))
            recit.append(Spacer(1, 4))
        elif type_bloc == "ol":
            for n, item in enumerate(contenu, 1):
                recit.append(Paragraph("%d. %s" % (n, enrichir(item)), styles["ul"]))
            recit.append(Spacer(1, 4))
        elif type_bloc == "quote":
            recit.append(Paragraph(enrichir(contenu), styles["quote"]))
        elif type_bloc == "code":
            for ligne in nettoyer(contenu, pour_pdf=True).split("\n"):
                recit.append(Paragraph(ligne.replace(" ", "&nbsp;") or "&nbsp;", styles["code"]))
        elif type_bloc == "hr":
            recit.append(Spacer(1, 6))
        elif type_bloc == "table" and contenu:
            nb_col = max(len(ligne) for ligne in contenu)
            largeur = 17 * cm
            poids = [max(1.2, min(6.0, 17.0 / nb_col))] * nb_col
            donnees = []
            for y, ligne in enumerate(contenu):
                rangee = []
                for x in range(nb_col):
                    texte = nettoyer(ligne[x] if x < len(ligne) else "", pour_pdf=True)
                    style = styles["cellh"] if y == 0 else styles["cell"]
                    rangee.append(Paragraph(enrichir(texte), style))
                donnees.append(rangee)
            tableau = PdfTable(donnees, colWidths=[largeur * p / sum(poids) for p in poids], repeatRows=1)
            tableau.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), violet),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd5e1")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
                ("LEFTPADDING", (0, 0), (-1, -1), 4), ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]))
            recit.append(KeepTogether(tableau))
            recit.append(Spacer(1, 8))

    def decor(canevas, doc):
        canevas.saveState()
        canevas.setFillColor(violet)
        canevas.rect(0, A4[1] - 0.6 * cm, A4[0], 0.6 * cm, stroke=0, fill=1)
        canevas.setFont("Helvetica", 7)
        canevas.setFillColor(gris)
        canevas.drawString(2 * cm, 1.1 * cm, "VOOMNET TECHNOLOGY - Gestion des achats")
        canevas.drawRightString(A4[0] - 2 * cm, 1.1 * cm, "Page %d" % doc.page)
        canevas.setStrokeColor(colors.HexColor("#e2e8f0"))
        canevas.line(2 * cm, 1.5 * cm, A4[0] - 2 * cm, 1.5 * cm)
        canevas.restoreState()

    modele = BaseDocTemplate(sortie, pagesize=A4, leftMargin=2 * cm, rightMargin=2 * cm,
                             topMargin=1.6 * cm, bottomMargin=1.8 * cm,
                             title=titre, author="VOOMNET TECHNOLOGY")
    cadre = Frame(2 * cm, 1.8 * cm, A4[0] - 4 * cm, A4[1] - 3.4 * cm, id="cadre")
    modele.addPageTemplates([PageTemplate(id="page", frames=[cadre], onPage=decor)])
    modele.build(recit)
    return sortie


# ------------------------------------------------------------------- main ---
if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("fichier")
    ap.add_argument("--titre", default=None)
    ap.add_argument("--sous-titre", default="VOOMNET TECHNOLOGY")
    args = ap.parse_args()

    chemin = args.fichier
    blocs = lire_markdown(chemin)
    titre = args.titre or next((nettoyer(c) for t, c in blocs if t == "h1"), "Document")
    base = os.path.splitext(chemin)[0]

    docx = generer_docx(blocs, base + ".docx", titre, args.sous_titre)
    pdf = generer_pdf(blocs, base + ".pdf", titre, args.sous_titre)
    for f in (docx, pdf):
        print("✔ %s (%.0f Ko)" % (os.path.basename(f), os.path.getsize(f) / 1024))
